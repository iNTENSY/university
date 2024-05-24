import os.path

from django.contrib import admin
from django.http import HttpResponse
from docx import Document

from university.settings import CURRENT_EDUCATION_HALF, BASE_DIR, CURRENT_DATE
from users.models import Groups
from .models import Cards, Attendance


@admin.register(Cards)
class CardsAdmin(admin.ModelAdmin):
    pass


@admin.register(Attendance)
class AttendanceAdmin(admin.ModelAdmin):
    pass


@admin.register(Groups)
class GroupsAdmin(admin.ModelAdmin):
    actions = ["download_attendance"]

    @admin.action(description="Получить данные по посещениям")
    def download_attendance(self, request, queryset):
        marked_groups = [group.title for group in queryset]
        statistic = [
            {
                "group": title,
                "people": Groups.attendance.make_for_group_by_title(
                    title,
                    CURRENT_EDUCATION_HALF["start"],
                    CURRENT_EDUCATION_HALF["end"]
                )
            }
            for title in marked_groups
        ]
        filepath = self.create_document(statistic)
        file = open(filepath, "rb")
        response = HttpResponse(file, content_type="application/msword")
        response["Content-Disposition"] = f'attachment; filename="{CURRENT_DATE}.doc"'
        file.close()
        os.remove(filepath)
        return response

    def create_document(self, statistic: list) -> str:
        document = Document()
        document.add_heading(f'Посещения студентов указанных групп c '
                             f'{CURRENT_EDUCATION_HALF["start"]} по {CURRENT_DATE}', 1)
        for group in statistic:
            document.add_paragraph()
            table = document.add_table(rows=1, cols=5)
            table.style = "Light Shading Accent 1"
            table.rows[0].cells[0].text = group["group"]
            for idx, data in enumerate(group["people"]):
                row = table.add_row().cells
                row[0].text = data.get("first_name", "")
                row[1].text = data.get("last_name", "")
                row[2].text = data.get("middle_name", "")
                row[3].text = str(data["counter"])
        filepath = os.path.join(BASE_DIR, f"templates/documents/{CURRENT_DATE}.doc")
        document.save(filepath)
        return filepath

